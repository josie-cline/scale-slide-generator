#!/usr/bin/env python3
"""
Generate a monthly status report by copying a .docx template and replacing
paragraph content for the current reporting period.

Usage:
    python3 generate_report.py
    python3 generate_report.py --template templates/my_template.docx
    python3 generate_report.py --output output/Feb_2026_MSR.docx

Customization:
    1. Place your MSR template in templates/
    2. Update TEMPLATE_PATH, OUTPUT_PATH, and PARAGRAPH_UPDATES below
    3. Or ask Cursor:
       "Update the MSR for reporting period 15 Feb - 14 Mar 2026 with these highlights: ..."
"""
import argparse
import shutil
from pathlib import Path

from docx import Document


# ═══════════════════════════════════════════════════════════════════════════
# CONFIGURATION — Edit for your program
# ═══════════════════════════════════════════════════════════════════════════

SCRIPT_DIR = Path(__file__).resolve().parent
TEMPLATE_PATH = SCRIPT_DIR / "templates" / "example_msr_template.docx"
OUTPUT_PATH = SCRIPT_DIR / "output" / "Monthly_Status_Report.docx"

# Paragraph replacements: paragraph_index -> new text.
# Run with --inspect to print all paragraph indices and their current text
# so you can identify which indices to update.
PARAGRAPH_UPDATES = {
    # Example: update title and reporting period
    # 0: "Scale AI - Monthly Report",
    # 3: "Reporting Period: 15 Feb - 14 Mar 2026",
}

# Table updates: (table_index, row_index) -> {col_index: new_text}
# Run with --inspect to see table structure.
TABLE_UPDATES = {
    # Example: update status column in use-case table
    # (1, 4): {3: "v1 Production", 7: "None."},
}


# ═══════════════════════════════════════════════════════════════════════════
# GENERATOR
# ═══════════════════════════════════════════════════════════════════════════

def inspect_template(template_path: Path):
    """Print all paragraphs and tables with indices for easy mapping."""
    doc = Document(str(template_path))

    print(f"\n{'='*60}")
    print(f"Template: {template_path.name}")
    print(f"{'='*60}\n")

    print("PARAGRAPHS:")
    print("-" * 40)
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(f"  [{i:3d}] {text[:100]}{'...' if len(text) > 100 else ''}")

    for t_idx, table in enumerate(doc.tables):
        print(f"\nTABLE {t_idx}: ({len(table.rows)} rows x {len(table.columns)} cols)")
        print("-" * 40)
        for r_idx, row in enumerate(table.rows):
            cells = [c.text.strip()[:30] for c in row.cells]
            print(f"  Row {r_idx}: {cells}")


def generate(template_path: Path, output_path: Path):
    output_path.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy(str(template_path), str(output_path))
    doc = Document(str(output_path))

    for idx, text in PARAGRAPH_UPDATES.items():
        if idx < len(doc.paragraphs):
            doc.paragraphs[idx].text = text

    for (t_idx, r_idx), col_updates in TABLE_UPDATES.items():
        if t_idx < len(doc.tables):
            table = doc.tables[t_idx]
            if r_idx < len(table.rows):
                for c_idx, text in col_updates.items():
                    if c_idx < len(table.rows[r_idx].cells):
                        table.rows[r_idx].cells[c_idx].text = text

    doc.save(str(output_path))
    print(f"Created {output_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate a monthly status report from a .docx template"
    )
    parser.add_argument(
        "--template", type=Path, default=TEMPLATE_PATH,
        help=f"Path to .docx template (default: {TEMPLATE_PATH})",
    )
    parser.add_argument(
        "--output", type=Path, default=OUTPUT_PATH,
        help=f"Output path (default: {OUTPUT_PATH})",
    )
    parser.add_argument(
        "--inspect", action="store_true",
        help="Print template structure (paragraphs and tables) without generating",
    )
    args = parser.parse_args()

    if args.inspect:
        inspect_template(args.template)
    else:
        generate(args.template, args.output)
