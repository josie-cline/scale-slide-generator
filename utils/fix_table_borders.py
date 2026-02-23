"""
Adds visible borders to all tables in the generated PoC Plan.
Run: python3 dla_ascend/fix_table_borders.py
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT = "dla_ascend/ASCEND_PoC_Plan_OP2.docx"
OUTPUT = "dla_ascend/ASCEND_PoC_Plan_OP2.docx"

def make_border_elm(tag, val="single", sz="4", color="1B1C1D"):
    el = OxmlElement(tag)
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)
    return el

def add_table_borders(table):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tblPr)

    # Remove existing tblBorders if present
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)

    tblBorders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tblBorders.append(make_border_elm(f"w:{side}"))
    tblPr.append(tblBorders)

    # Also set 100% table width
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "5000")
    tblW.set(qn("w:type"), "pct")
    tblPr.append(tblW)

    # Style header row cells with light grey shading
    if table.rows:
        hdr_row = table.rows[0]
        for cell in hdr_row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "E8E8E8")
            for existing in tcPr.findall(qn("w:shd")):
                tcPr.remove(existing)
            tcPr.append(shd)

def main():
    doc = Document(INPUT)
    for table in doc.tables:
        add_table_borders(table)
    doc.save(OUTPUT)
    print(f"Tables fixed and saved to {OUTPUT}")

if __name__ == "__main__":
    main()
