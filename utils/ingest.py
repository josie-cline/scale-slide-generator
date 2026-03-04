"""
Document-to-DECK ingestion: parse .pdf, .docx, .csv, .txt, .json into
structured data suitable for building a DECK dict.

Usage:
    from utils.ingest import ingest_file
    result = ingest_file("path/to/file.csv")
    # result = {"type": "csv", "headers": [...], "rows": [...], "raw_text": "..."}
"""
import csv
import json
from pathlib import Path


def ingest_file(path: str) -> dict:
    """Parse a file and return structured content.

    Returns a dict with keys:
        type     — file type (csv, json, txt, docx, pdf)
        raw_text — full text content
        headers  — column headers (csv/table data)
        rows     — row data (csv/table data)
        data     — parsed structure (json files)
        error    — error message if parsing failed
    """
    p = Path(path)
    if not p.exists():
        return {"type": "unknown", "error": f"File not found: {path}"}

    suffix = p.suffix.lower()
    try:
        if suffix == ".csv":
            return _ingest_csv(p)
        elif suffix == ".json":
            return _ingest_json(p)
        elif suffix == ".txt" or suffix == ".md":
            return _ingest_text(p)
        elif suffix == ".docx":
            return _ingest_docx(p)
        elif suffix == ".pdf":
            return _ingest_pdf(p)
        else:
            return _ingest_text(p)
    except Exception as e:
        return {"type": suffix.lstrip("."), "error": str(e)}


def _ingest_csv(p: Path) -> dict:
    with open(p, newline="", encoding="utf-8-sig") as f:
        reader = csv.reader(f)
        all_rows = list(reader)

    if not all_rows:
        return {"type": "csv", "headers": [], "rows": [], "raw_text": ""}

    headers = all_rows[0]
    rows = all_rows[1:]
    raw_text = "\n".join(",".join(row) for row in all_rows)
    return {"type": "csv", "headers": headers, "rows": rows, "raw_text": raw_text}


def _ingest_json(p: Path) -> dict:
    raw = p.read_text(encoding="utf-8")
    data = json.loads(raw)

    result = {"type": "json", "data": data, "raw_text": raw}

    if isinstance(data, list) and data and isinstance(data[0], dict):
        headers = list(data[0].keys())
        rows = [[str(row.get(h, "")) for h in headers] for row in data]
        result["headers"] = headers
        result["rows"] = rows

    return result


def _ingest_text(p: Path) -> dict:
    text = p.read_text(encoding="utf-8")
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return {"type": "txt", "raw_text": text, "lines": lines}


def _ingest_docx(p: Path) -> dict:
    try:
        from docx import Document
    except ImportError:
        return {"type": "docx", "error": "python-docx not installed"}

    doc = Document(str(p))
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    raw_text = "\n".join(paragraphs)

    tables = []
    for tbl in doc.tables:
        rows = []
        for row in tbl.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append({"headers": rows[0], "rows": rows[1:]})

    return {
        "type": "docx",
        "raw_text": raw_text,
        "paragraphs": paragraphs,
        "tables": tables,
    }


def _ingest_pdf(p: Path) -> dict:
    try:
        import subprocess
        result = subprocess.run(
            ["python3", "-c",
             f"import pdfplumber; pdf = pdfplumber.open('{p}'); "
             f"print('\\n'.join(page.extract_text() or '' for page in pdf.pages))"],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode == 0:
            text = result.stdout
            lines = [l.strip() for l in text.splitlines() if l.strip()]
            return {"type": "pdf", "raw_text": text, "lines": lines}
    except Exception:
        pass

    return {"type": "pdf", "error": "pdfplumber not available — install with: pip install pdfplumber"}


def ingest_to_slides(path: str) -> "list[dict]":
    """Best-effort conversion of a file into slide dicts.

    CSV/JSON with tabular data → table slide.
    Text/docx → content slides (one per ~5 paragraphs).
    """
    result = ingest_file(path)

    if result.get("error"):
        return [{"layout": "content", "title": "Import Error", "bullets": [result["error"]]}]

    slides = []
    filename = Path(path).stem.replace("_", " ").title()

    if result.get("headers") and result.get("rows"):
        slides.append({
            "layout": "table",
            "title": filename,
            "headers": result["headers"],
            "rows": result["rows"][:15],
        })

    if result["type"] == "docx" and result.get("tables"):
        for tbl in result["tables"]:
            slides.append({
                "layout": "table",
                "title": filename,
                "headers": tbl["headers"],
                "rows": tbl["rows"][:15],
            })

    lines = result.get("lines") or result.get("paragraphs") or []
    if lines and not slides:
        chunk_size = 5
        for i in range(0, len(lines), chunk_size):
            chunk = lines[i:i + chunk_size]
            title = chunk[0] if len(chunk[0]) < 80 else filename
            bullets = chunk[1:] if title == chunk[0] else chunk
            slides.append({
                "layout": "content",
                "title": title,
                "bullets": bullets[:8],
            })

    if not slides:
        slides.append({"layout": "content", "title": filename, "bullets": ["No extractable content found"]})

    return slides
