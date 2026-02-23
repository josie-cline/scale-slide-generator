"""
Generates the updated ASCEND PoC Plan for Option Period 2.
Run from workspace root: python3 dla_ascend/generate_poc_plan_op2.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy
import os

TEMPLATE_PATH = "_inbox/Prototype Proof of Concept Plan - OY1 .docx"
OUTPUT_PATH = "dla_ascend/ASCEND_PoC_Plan_OP2.docx"

def add_heading(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)
    return p

def add_body(doc, text, bold=False):
    p = doc.add_paragraph(style="normal")
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)
    if bold:
        run.font.bold = True
    return p

def add_table_row(table, cells, bold_first=False):
    row = table.add_row()
    for i, val in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)
        if bold_first and i == 0:
            run.font.bold = True
    return row

def build_doc():
    # Start from scratch but clone styles from template
    template = Document(TEMPLATE_PATH)
    doc = Document(TEMPLATE_PATH)

    # Clear all content while preserving styles and section properties
    body = doc.element.body
    # Keep the last sectPr element (page layout), remove everything else
    for element in list(body)[:-1]:
        body.remove(element)

    # ── TITLE BLOCK ───────────────────────────────────────────────
    p = doc.add_paragraph(style="Heading 3")
    run = p.add_run("Prototype Proof of Concept Plan: Option Period 2")
    run.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)
    run.font.bold = True

    p = doc.add_paragraph(style="normal")
    run = p.add_run("Plan in support of Project Ascend")
    run.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)
    run.font.bold = True

    # ── HEADER TABLE ──────────────────────────────────────────────
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = doc.styles["TableNormal"]
    header_data = [
        ("Contract Number:", "HC1084-25-0001"),
        ("Participant Name:", "Scale AI"),
        ("AOR:", "Melissa Redd, melissa.j.redd.civ@mail.mil"),
        ("Participant POC:", "Josiah Cline, josiah.cline@scale.com"),
        ("Date:", "February 20, 2026"),
    ]
    # Use first row for first item, add rows for rest
    tbl.rows[0].cells[0].text = ""
    tbl.rows[0].cells[1].text = ""
    r0 = tbl.rows[0]
    run0 = r0.cells[0].paragraphs[0].add_run(header_data[0][0])
    run0.font.bold = True
    run0.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)
    run1 = r0.cells[1].paragraphs[0].add_run(header_data[0][1])
    run1.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)
    for label, val in header_data[1:]:
        add_table_row(tbl, [label, val], bold_first=True)

    doc.add_paragraph()

    # ── PURPOSE ───────────────────────────────────────────────────
    add_heading(doc, "Purpose")
    add_body(doc,
        "The objective for Option Period 2 is to sustain and mature the Donovan platform capabilities "
        "delivered during the Base Period and Option Period 1, while completing outstanding deliverables "
        "and refining existing applications based on user feedback. The goal is to ensure continued "
        "operational availability of all nine (9) workflow applications on SIPR, complete the live "
        "ServiceNow data integration, and advance two priority applications from v0 to v1 in support "
        "of DISA's Defensive Cyber Operations (DCO) mission."
    )

    # ── SCOPE ─────────────────────────────────────────────────────
    add_heading(doc, "Scope")
    add_body(doc,
        "This plan covers the 3-month Option Period 2 from February 15, 2026, to May 14, 2026. "
        "The scope for this period includes:"
    )
    scope_items = [
        "Platform Sustainment: Continued platform maintenance, updates, and bug fixes for all nine (9) "
        "applications delivered during the Base Period and Option Period 1, for up to 200 Donovan accounts on SIPR.",
        "Application Refinement (v0 → v1): Advancing the INTEL Agent (App 9) and Incident Response Review "
        "(App 5) from v0 to v1 in accordance with Government-provided feedback, no later than 14 days "
        "following contract award.",
        "Live Data Integration: Completing the one (1) ServiceNow data connection identified under the "
        "Option Period 1 agreement.",
        "ATO Maintenance: Providing all required documentation and artifacts necessary to maintain the "
        "Authority to Operate (ATO-C), which is active through August 2026.",
        "Optional Task 1 (CLIN 0013): Development of up to three (3) new workflow applications of "
        "comparable size, scope, and complexity as the existing nine, if exercised.",
        "Optional Task 2 (CLIN 0014): Development of new live data connections to existing Government "
        "datastores, if exercised.",
    ]
    for item in scope_items:
        p = doc.add_paragraph(style="normal")
        run = p.add_run(item)
        run.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)
        p.paragraph_format.left_indent = Pt(18)

    # ── SOLUTION OVERVIEW ─────────────────────────────────────────
    add_heading(doc, "Solution Overview")
    add_body(doc,
        "Scale's Donovan platform continues to serve as a force multiplier for DISA's Global DODIN "
        "workforce. During the Base Period and Option Period 1, Scale deployed Donovan to DISA's SIPRNet "
        "environment, connected Mission Group datasets, and delivered nine (9) custom RAG-enabled workflow "
        "applications aligned to DCO use cases. The platform is now operational with an active ATO-C "
        "(valid through August 6, 2026) and supporting up to 200 users across DISA Mission Groups."
    )
    add_body(doc,
        "Option Period 2 shifts focus from delivery to sustainment and refinement. Scale will maintain "
        "platform stability, complete the ServiceNow integration, mature two priority applications to v1, "
        "and continue forward-deployed support to maximize user adoption and mission impact."
    )

    # ── BASE PERIOD COMPLETED APPS ────────────────────────────────
    add_heading(doc, "Base Period: Completed Workflow Applications (Apps 1–3)")
    add_body(doc,
        "The following three workflow applications were developed and delivered during the Base Period "
        "on NIPRNet. These applications are available on SIPR pending ATO-C cutover and are sustained "
        "under Option Period 2."
    )
    base_apps = [
        ("App 1 — Threat Cards",
         "Provides RAG-based prototype functionality to search across user-selected knowledge bases "
         "composed of text-based documents, synthesizing discovered information into a comprehensive, "
         "pre-formatted Threat Card with source citations. Designed as an ~80% solution to reduce manual "
         "research and data synthesis burden."),
        ("App 2 — Cyber Highlights",
         "Processes single text-based PDF or .txt files uploaded by a user, delivering extracted "
         "key cyber-related information in a summarized format. The resulting summary reduces the time "
         "required for manual data compilation, enabling analysts to focus on final analysis and refinement."),
        ("App 3 — Splunk Extraction",
         "Processes a single text-based PDF or .txt file and delivers all relevant technical details "
         "in a structured format ready for Splunk ingestion. Eliminates manual data entry, freeing "
         "analysts from repetitive compilation tasks."),
    ]
    for title, desc in base_apps:
        p = doc.add_paragraph(style="normal")
        run = p.add_run(title)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)
        add_body(doc, desc)

    # ── OPTION PERIOD 1: COMPLETED APPS ───────────────────────────
    add_heading(doc, "Option Period 1: Delivered Workflow Applications (Apps 4–9)")
    add_body(doc,
        "All six (6) mutually agreed-upon custom workflow applications were delivered on SIPRNet during "
        "Option Period 1 (August 15, 2025 – February 14, 2026). The following applications are "
        "operational and sustained under Option Period 2."
    )
    op1_apps = [
        ("App 4 — RFI Response (Threat Hunt Planning)",
         "RAG-based prototype that searches intelligence reports and documents to create summarized "
         "drafts for quick RFI responses with citations to source materials. Enables rapid construction "
         "of tailored RFI responses and replaces manual from-scratch data discovery. "
         "Delivered: October 2025."),
        ("App 5 — Incident Response Review (v0 → v1 in OP2)",
         "RAG-based tool that processes Incident Response (IR) tippers and tickets in CSV format to "
         "extract and summarize key threat indicators and actionable intelligence for hunt teams. "
         "Advancing from v0 to v1 in Option Period 2 based on Government feedback. "
         "Delivered: October 2025."),
        ("App 6 — Historical Retro Analysis",
         "RAG-based capability enabling analysts to perform natural language queries against knowledge "
         "bases containing IRs, PULSE, ACE, and SIGACTS. Retrieves relevant information across "
         "disparate historical datasets and reduces time required to locate specific historical "
         "intelligence. Delivered: October 2025."),
        ("App 7 — CCIR Decision Support System",
         "RAG-enabled system that analyzes incident reports against Commander's Critical Information "
         "Requirements (CCIR) criteria, providing instant analysis and pre-formatted notifications. "
         "Knowledge base includes historical incident logs and CCIR matrices. "
         "Delivered: December 2025."),
        ("App 8 — Mission Assurance Site Risk Analysis",
         "RAG-enabled knowledge base that analyzes historical incident reports to rank sites based on "
         "risk, providing a prioritized list and heat map visualization. Outputs machine-readable formats "
         "(CSV, XML, JSON) for integration with mission assurance planning tools. "
         "Delivered: December 2025."),
        ("App 9 — Intel Agent (v0 → v1 in OP2)",
         "RAG-enabled system that extracts actionable intelligence from SIGACTs and TIPPER threat "
         "reports, including Indicators of Compromise (IOCs), TTPs, and recommended mitigations, "
         "formatted for direct ingestion into security tools such as Splunk and Elastic. "
         "Advancing from v0 to v1 in Option Period 2 based on Government feedback. "
         "Delivered: December 2025."),
    ]
    for title, desc in op1_apps:
        p = doc.add_paragraph(style="normal")
        run = p.add_run(title)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)
        add_body(doc, desc)

    # ── EXPECTED OUTCOMES & MILESTONES ────────────────────────────
    add_heading(doc, "Expected Outcomes & Milestones")
    add_body(doc,
        "The work in Option Period 2 is structured around four key tasks designed to sustain platform "
        "operations, complete outstanding deliverables, and mature priority applications."
    )

    tasks = [
        ("Task 1: Option Period 2 Kickoff (PWS 6.4.1)",
         "Scale will confirm objectives for Option Period 2, including the v0 → v1 refinement scope "
         "for Apps 5 and 9, the plan to complete the ServiceNow data connection, and success criteria "
         "for the option period. The key deliverable for this task is the Option Period 2 Prototype "
         "Proof of Concept Plan."),
        ("Task 2: Platform Sustainment and Application Refinement (PWS CLIN 0004/0008)",
         "Scale will provide continuous Donovan platform maintenance, updates, and bug fixes for all "
         "nine (9) applications on Government-hosted SIPR for up to 200 users. Within 14 days of "
         "contract award, Apps 5 (IR Review) and 9 (Intel Agent) will advance from v0 to v1 in "
         "accordance with Government-provided feedback. Scale will also complete the ServiceNow live "
         "data connection upon receipt of the required Government POC approval and access."),
        ("Task 3: Optional — New Application Development (CLIN 0013)",
         "If exercised, Scale will develop up to three (3) new workflow applications of comparable "
         "size, scope, and complexity as the existing nine applications. Scope and technical details "
         "for new applications will be mutually defined through Technical Exchange Meetings (TEMs) "
         "prior to development."),
        ("Task 4: Optional — New Data Connections (CLIN 0014)",
         "If exercised, Scale will develop new live data connections to existing Government datastores "
         "as mutually agreed upon. Integration scope and technical feasibility will be confirmed "
         "through TEMs prior to engineering work."),
    ]
    for title, desc in tasks:
        p = doc.add_paragraph(style="normal")
        run = p.add_run(title)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)
        add_body(doc, desc)

    # ── DELIVERABLES TABLE ────────────────────────────────────────
    add_heading(doc, "Deliverables")
    del_tbl = doc.add_table(rows=1, cols=3)
    del_tbl.style = doc.styles["TableNormal"]
    hdr = del_tbl.rows[0]
    for i, h in enumerate(["Task #", "Deliverable", "Due Date"]):
        hdr.cells[i].text = ""
        run = hdr.cells[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)

    deliverables = [
        ("CLIN 0004/0008", "Platform Sustainment — Donovan licenses, maintenance, and bug fixes for 9 apps (200 users)", "Duration of OP2 (Feb 15 – May 14, 2026)"),
        ("CLIN 0004/0008", "App v0 → v1 Refinements: Intel Agent (App 9) and IR Review (App 5)", "NLT 14 days post OP2 award"),
        ("CLIN 0004/0008", "ServiceNow Live Data Integration (1 connection)", "Upon receipt of Government POC approval"),
        ("CLIN 0004/0008", "ATO Documentation and Artifacts", "As required"),
        ("CLIN 0013 (Optional)", "Up to three (3) new workflow applications", "Mutually agreed upon, per TEM"),
        ("CLIN 0014 (Optional)", "New data connections to Government datastores", "Mutually agreed upon, per TEM"),
    ]
    for row_data in deliverables:
        add_table_row(del_tbl, list(row_data))

    doc.add_paragraph()

    # ── TIMELINE ─────────────────────────────────────────────────
    add_heading(doc, "Timeline")
    add_body(doc, "Option Period 2 Period of Performance: February 15, 2026 – May 14, 2026 (3 months)")
    doc.add_paragraph()

    tl_tbl = doc.add_table(rows=1, cols=3)
    tl_tbl.style = doc.styles["TableNormal"]
    hdr2 = tl_tbl.rows[0]
    for i, h in enumerate(["Milestone", "Target Date", "Status"]):
        hdr2.cells[i].text = ""
        run = hdr2.cells[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)

    timeline_rows = [
        ("OP2 Kickoff Meeting", "February 20, 2026", "Complete"),
        ("OP2 PoC Plan Delivered", "February 20, 2026", "Complete"),
        ("Apps 5 & 9: v0 → v1 Refinements", "NLT 14 days post award", "In Progress"),
        ("ServiceNow Integration Complete", "Pending Government POC approval", "In Progress"),
        ("Optional Tasks (CLIN 0013/0014)", "Mutually agreed upon", "TBD if exercised"),
        ("End of Option Period 2 PoP", "May 14, 2026", "Upcoming"),
    ]
    for row_data in timeline_rows:
        add_table_row(tl_tbl, list(row_data))

    doc.add_paragraph()

    # Add context for prior periods
    add_body(doc, "For reference, the full program timeline is as follows:")
    full_tl = doc.add_table(rows=1, cols=3)
    full_tl.style = doc.styles["TableNormal"]
    hdr3 = full_tl.rows[0]
    for i, h in enumerate(["Period", "Dates", "Status"]):
        hdr3.cells[i].text = ""
        run = hdr3.cells[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)

    full_periods = [
        ("Base Period", "November 15, 2024 – August 14, 2025", "Complete"),
        ("Option Period 1", "August 15, 2025 – February 14, 2026", "Complete"),
        ("Option Period 2 (CLIN 0004/0008)", "February 15, 2026 – May 14, 2026", "Active"),
        ("Option Period 3 (CLIN 0009/0010)", "May 15, 2026 – August 14, 2026", "Optional"),
        ("Option Period 4 (CLIN 0011/0012)", "August 15, 2026 – November 14, 2026", "Optional"),
    ]
    for row_data in full_periods:
        add_table_row(full_tl, list(row_data))

    doc.add_paragraph()

    # ── RISKS & DEPENDENCIES ──────────────────────────────────────
    add_heading(doc, "Risks & Dependencies")
    add_body(doc,
        "The risks for Option Period 2 are focused on integration completion and application refinement. "
        "The following table identifies key dependencies and corresponding mitigations."
    )
    doc.add_paragraph()

    risk_tbl = doc.add_table(rows=1, cols=2)
    risk_tbl.style = doc.styles["TableNormal"]
    hdr4 = risk_tbl.rows[0]
    for i, h in enumerate(["Risk / Dependency", "Mitigation"]):
        hdr4.cells[i].text = ""
        run = hdr4.cells[i].paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)

    risks = [
        (
            "ServiceNow Integration — Government POC Approval Required: "
            "Completion of the ServiceNow data connection is dependent on receipt of Government POC designation "
            "and approval to configure the Donovan <> ServiceNow API.",
            "Scale has submitted architecture diagrams and API endpoints are ready for testing on ScaleGov. "
            "Scale will proceed with integration testing immediately upon Government approval. "
            "A Technical Exchange Meeting (TEM) will be scheduled upon POC designation to de-risk the timeline."
        ),
        (
            "Application v0 → v1 Scope Definition: The specific feedback and refinement requirements "
            "for Apps 5 and 9 must be provided by the Government to guide v1 development.",
            "Scale will request written feedback documentation within the first week of OP2 to ensure "
            "a 14-day delivery timeline can be met. If feedback is not received within 5 business days "
            "of contract award, Scale will schedule a TEM to capture requirements verbally."
        ),
        (
            "Optional CLIN Scope Definition (CLIN 0013/0014): If optional tasks are exercised, "
            "timely agreement on new application scope and data connection requirements is necessary "
            "to meet the OP2 period of performance.",
            "Scale recommends initiating Joint Application Design (JAD) sessions within the first "
            "month of OP2 to define and approve new application or data connection specifications, "
            "should the Government elect to exercise these CLINs."
        ),
        (
            "Infrastructure Performance: Continued user growth and live data queries may increase "
            "computational load on the SIPR deployment.",
            "Scale will conduct performance monitoring and stress testing as new capabilities are "
            "activated. Scale will coordinate with the appropriate Government infrastructure representatives "
            "to provision additional resources if performance degradation is identified."
        ),
    ]
    for risk, mitigation in risks:
        row = risk_tbl.add_row()
        for i, text in enumerate([risk, mitigation]):
            row.cells[i].text = ""
            run = row.cells[i].paragraphs[0].add_run(text)
            run.font.color.rgb = RGBColor(0x1B, 0x1C, 0x1D)

    # Save
    os.makedirs("dla_ascend", exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")

if __name__ == "__main__":
    build_doc()
