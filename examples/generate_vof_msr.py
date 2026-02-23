#!/usr/bin/env python3
"""
Generate Valley of Fire MSR by copying template and replacing content in place.
Preserves template format 1:1 - only updates content for the reporting period.
Output: VoF_Monthly_Status_Report_16Jan-14Feb_2026.docx (overwrites)
"""

import shutil
from pathlib import Path

from docx import Document

_SCRIPT_DIR = Path(__file__).resolve().parent
_PROJECT_ROOT = _SCRIPT_DIR.parent
TEMPLATE = _PROJECT_ROOT / 'scale_ai_disa' / 'ScaleAi DISA - JAN 2026 MSR.docx'
OUTPUT = _SCRIPT_DIR / 'VoF_Monthly_Status_Report_16Jan-14Feb_2026.docx'


def main():
    shutil.copy(str(TEMPLATE), str(OUTPUT))
    doc = Document(str(OUTPUT))

    # Paragraph replacements - only paragraphs with content; preserve structure
    # Index -> new text. Empty/whitespace paras (8, 23, 82, etc) stay as-is.
    para_updates = {
        0: 'Scale AI - Monthly Report',
        1: 'Contract: HC1084259001',
        2: 'Total Period of Performance: 15 AUG 2025 to 14 FEB 2026',
        3: 'Reporting Period: 16 JAN - 14 FEB',
        6: 'Executive Summary ',
        9: 'Infrastructure: Architecture diagrams for ServiceNow integration submitted to VoF leadership. API endpoints ready on ScaleGov. Connection blocked pending DISA CCB approval.',
        11: 'Platform: Donovan stable on SIPR. ATO-C active through August 2026. Monthly update schedule maintained.',
        13: 'Use Cases: All six OP1 use cases delivered. Four apps in v1 Production (Threat Hunting Intel, Historical Analysis, CCIR Decision Support, Site Risk Analysis). Two apps in v0 awaiting demos (IR Review, INTEL Agent). End-user testing and validation pending.',
        14: 'Security/Compliance: ATO-C active through August 2026. SIPR credentials expire 14 February 2026. OP2 renewal and access continuity pending USG confirmation.',
        24: 'Contract: NSTR',
        34: 'Engagement',
        36: 'Monthly Status Report Rollup covering the final month of Valley of Fire Option Period 1.',
        37: 'Monthly Status Report Rollup: 16 January 2026 – 14 February 2026',
        38: '1. Executive Summary',
        39: 'January through mid-February focused on OP1 close-out, v0 to v1 refinements, and joint demonstration preparation. All six custom workflow applications delivered to DISA-PAC and DISA-EUR. Four apps in v1 Production. End-user testing and validation pending. ServiceNow integration blocked pending DISA CCB. Contract mod response completed 3 February. OP1 PoP concludes 14 February 2026.',
        40: '2. Use Case Development & Demonstrations',
        41: 'Threat Hunt (TH) App (App 4):',
        42: 'Status: v1 Production — DISA-PAC. Natural language queries over RAG-enabled knowledge bases for RFI response generation. End-user validation pending.',
        43: 'Feedback: Production deployment complete. Citations and RAG-enabled responses operational.',
        44: 'Requirements: Extracting IOCs, specific malware/APT data, attack chain info — addressed in production.',
        45: 'Action Items: None outstanding.',
        46: 'Site Assurance (Risk Analysis) App (App 8):',
        47: 'Status: v1 Production — DISA-EUR. Automated site risk ranking; CSV export for mission assurance.',
        48: 'Status: Defined process and data sources. Production deployment complete.',
        49: 'Action Items: None outstanding.',
        50: 'IR Review (App 5):',
        51: 'Status: v0 Awaiting Demo — DISA Global. Tipper/ticket CSV analysis for incident data. End-user validation pending.',
        52: 'Action Items: Schedule demo; v0 to v1.',
        53: 'Historical Retrospective (App 6): Status: v1 Production — DISA-PAC. Unified cross-source search. Action Items: None outstanding.',
        54: 'CCIR App (App 7):',
        55: 'Status: v1 Production — DISA-EUR. Demo conducted and resonated with team.',
        56: 'Workflow: Incident text → search KB for similar incidents → CCIR Matrix → MET/NOT MET determination.',
        57: 'Action Items: None outstanding.',
        58: 'INTEL Agent (App 9):',
        59: 'Status: v0 Awaiting Demo — DISA-EUR. IOC extraction and Splunk query generation from threat reports. End-user validation pending.',
        60: 'Action Items: Schedule demo; v0 to v1.',
        61: '3. Program Management & Stakeholder Engagement',
        62: 'Schedule & Contract:',
        63: 'Contract Mod Response: Completed per 3 February due date.',
        64: 'OP1 Close-Out: Documentation in progress per Task 3. Joint demo event in preparation.',
        65: 'OP2 Planning: Renewal pending USG confirmation. Re-credentialing adds 1-2 months if OP2 proceeds.',
        66: 'Stakeholder Coordination:',
        67: 'ServiceNow: Deepak steering CCB approval. Architecture submitted; Thomas Walsh ready for endpoint validation.',
        68: 'VoF Actions Required: Monitor ServiceNow CCB, itemize app updates for PoP close, confirm renewal.',
        69: 'Threat Card: Analyst engagement as scheduling permits during OP1 close-out.',
        70: '4. Technical Infrastructure & Challenges',
        71: 'SIPR Tokens & Onboarding:',
        72: 'SIPR credentials (CAC/SIPR) expire 14 February 2026, coinciding with OP1 PoP end.',
        73: 'Tom: Token acquired. Nisha: Per prior report.',
        74: 'Next Steps: Coordinate re-credentialing if OP2 proceeds.',
        75: 'Data & Integrations:',
        76: 'ServiceNow: Blocked. Awaiting DISA CCB approval. Upon approval: automated incident ingestion, real-time sync for RAG apps.',
        77: 'Known Bugs & Performance (Timeout/Uploads):',
        78: 'Issue: Timeout and upload issues addressed in prior period.',
        79: 'Resolution Status: Fix deployed Jan. Platform stability maintained.',
        80: 'Verification: OP1 close-out verification in progress.',
        81: 'Outstanding Action Items: ServiceNow CCB approval; App 5 and 9 demos; v0 to v1 refinements post-demo.',
        87: 'Use Case Table',
        91: 'Next Month',
        92: 'In order of priority, Engagement:',
        93: 'Complete OP1 Close-Out: Finalize documentation, joint demo event, Task 3 deliverables.',
        94: 'Finalize v0-v1 Apps: Complete IR Review and INTEL Agent demos; incorporate feedback; promote to v1.',
        95: 'Finalize Contracting Misalignment (Scale <> DISA): Resolve outstanding contractual items, close out and invoice CLIN 0003/0005/0007; align on OP2 scope.',
        96: 'Finalize ServiceNow Data Connection: Establish connection upon DISA CCB approval.',
        97: 'Support End User Onboarding/Training: Facilitate training sessions and workflow integration for identified user groups.',
        98: 'In order of priority, Infrastructure:',
        99: 'Credential Transition: Coordinate re-credentialing if OP2; plan 1-2 month lead time.',
        100: 'In order of priority, Compliance:',
        101: 'ATO-C: Maintain through August 2026. No motions required for OP1 close.',
    }

    for i, para in enumerate(doc.paragraphs):
        if i in para_updates:
            para.text = para_updates[i]

    # Table 0 (Tasks): Keep exactly as template - no changes
    # Table 1 (Use Case): Only update Status (C3) and What are we going to do (C7) for rows 4-9
    t1 = doc.tables[1]
    t1_updates = {
        4: ('v1 Production', 'None.'),
        5: ('v0 Awaiting Demo', 'Schedule demo; v0 to v1.'),
        6: ('v1 Production', 'None.'),
        7: ('v1 Production', ''),
        8: ('v1 Production', ''),
        9: ('v0 Awaiting Demo', 'Schedule demo; v0 to v1.'),
    }
    for row_idx, (status, what) in t1_updates.items():
        t1.rows[row_idx].cells[3].text = status
        t1.rows[row_idx].cells[7].text = what

    # Table 2 (Financials): Keep exactly as template

    doc.save(str(OUTPUT))
    print(f'Updated: {OUTPUT}')


if __name__ == '__main__':
    main()
